from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\clasp\Documents\GitHub\Bakery-Audit\bot_new_features_report.docx"

ACCENT = RGBColor(0xC8, 0x5C, 0x3D)
INK = RGBColor(0x20, 0x1B, 0x18)
MUTED = RGBColor(0x6B, 0x63, 0x5C)


def add_rule(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C85C3D")
    pbdr.append(bottom)
    p_pr.append(pbdr)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(9.5)
normal.font.color.rgb = INK
pf = normal.paragraph_format
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.line_spacing = 1.07

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = title.add_run("Bakery-Audit: New Bot Features")
r.bold = True
r.font.name = "Georgia"
r.font.size = Pt(20)
r.font.color.rgb = INK

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(4)
r = sub.add_run("Strict one-page implementation report. Scope: newly added odds-analysis and date-resolution features only.")
r.italic = True
r.font.size = Pt(9.2)
r.font.color.rgb = MUTED

rule = doc.add_paragraph()
rule.add_run(" ")
add_rule(rule)
rule.paragraph_format.space_after = Pt(6)

sections = [
    (
        "Odds automation",
        "Added a dedicated `@bot odds [real|bonus|both]` flow. Screenshots are parsed into `OddsCandidate` rows, reviewed in Discord, then written to `odds_raw`, `odds_clean`, and `odds_ranked` worksheet tabs.",
    ),
    (
        "Moneyline",
        "Moneyline now supports cross-site bet/hedge pairing and workbook-style calculations. The pipeline groups opposite sides of the same matchup by date, rejects same-site pairs, applies the Cloudbet bonus-side restriction, and ranks the best reportable direction.",
    ),
    (
        "Over / Under",
        "Totals are extracted as `total_over` / `total_under` plus `total_line`. The bot pairs Over and Under across sites for the same game, optimizes hedge size with `_optimize_ou_hedge_stake(...)`, and evaluates worst-case and middle-window score outcomes.",
    ),
    (
        "Spread",
        "Spread bets are extracted as `market=\"spread\"` plus signed `spread_line`. The bot pairs opposite spread sides across sites, optimizes hedge size with `_optimize_spread_hedge_stake(...)`, and evaluates margin buckets around the line thresholds.",
    ),
    (
        "Ranking and output",
        "Recommendation selection now avoids duplicate games across metric buckets, preserves sportsbook attribution more aggressively, aligns rake / edge sign to the workbook display, and replaces the older multi-page odds results with one best pick each for moneyline, O/U, and spread.",
    ),
    (
        "Date handling for testing",
        "The bot now reads image capture metadata when available and uses that date as the reference for relative labels. `today`, `tomorrow`, and `yesterday` resolve against the image date rather than the server clock, which makes historical screenshot testing reliable. Month/day values without a year now inherit the reference year.",
    ),
]

for heading, body in sections:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(heading)
    r.bold = True
    r.font.size = Pt(10.3)
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(2)
    lead = p.add_run("- ")
    lead.bold = True
    lead.font.color.rgb = ACCENT
    r = p.add_run(body)
    r.font.size = Pt(9.4)
    r.font.color.rgb = INK

footer = doc.add_paragraph()
footer.paragraph_format.space_before = Pt(5)
r = footer.add_run(
    "Implementation anchors: src/bot/odds_pipeline.py, src/bot/odds_ui.py, src/bot/gemini_client.py, src/bot/models.py, src/bot/image_metadata.py."
)
r.font.size = Pt(8.3)
r.font.color.rgb = MUTED

doc.save(OUTPUT)
print(OUTPUT)
